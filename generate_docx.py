# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def md_to_docx(md_file, docx_file):
    doc = Document()
    
    encodings = ['utf-8-sig', 'utf-8', 'gbk', 'latin-1']
    content = None
    for enc in encodings:
        try:
            with open(md_file, 'r', encoding=enc) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue
    
    if content is None:
        raise ValueError(f"无法读取文件 {md_file}")
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                if code_content:
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.style = doc.styles['Normal']
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        if line.startswith('# '):
            heading = line[2:].strip()
            doc.add_heading(heading, level=1)
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
        elif line.startswith('---'):
            pass
        elif line.startswith('| ') and '|' in line[2:]:
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                if not lines[i].startswith('|--'):
                    table_lines.append(lines[i])
                i += 1
            if table_lines:
                create_table(doc, table_lines)
            continue
        elif line.strip():
            clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
            clean_line = re.sub(r'\*([^*]+)\*', r'\1', clean_line)
            clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)
            clean_line = ''.join(char for char in clean_line if ord(char) >= 32 or char in '\n\r\t')
            if clean_line.strip():
                p = doc.add_paragraph(clean_line)
        
        i += 1
    
    doc.save(docx_file)
    print(f'✅ {docx_file} 生成成功')

def create_table(doc, lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                row.cells[j].text = cell_text

if __name__ == '__main__':
    md_to_docx('skill.md', 'skill.docx')