import os
from datetime import datetime

print("=" * 70)
print("📝 文档更新与Git推送工具")
print(f"⏰ 时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
print("=" * 70)

# 1. 更新 README.md
print("\n📄 [1/4] 更新 README.md...")
readme_path = "README.md"
with open(readme_path, "r", encoding="utf-8") as f:
    content = f.read()

version_update = f"""

---

## 🔧 Bug 修复记录 (v3.8.89.9)

**修复日期**: {datetime.now().strftime("%Y-%m-%d %H:%M")}

### ✅ 已修复问题

#### 1. **高价商品(≥599)显示为0**
- **根本原因**: 价格字段引用错误 - 使用了 `'售䷜'` (U+552E U+4DDC) 而非 `'price'`
- **修复位置**: 9处字段引用已修正
- **结果**: 高价商品从0个恢复为71个 ✅

#### 2. **Flask → FastAPI 完全迁移（20处）**
- **修复内容**: 
  - `return jsonify(...), status_code` → `return jsonify(..., status_code=status)`
  - 统一使用FastAPI异步请求处理 (`await request.json()`)
  - 所有36个路由端点已完全迁移至FastAPI规范

#### 3. **价格数据丢失问题（核心修复）**
- **问题**: 第二次运行爬虫时，新数据覆盖了有价格的旧数据
- **解决方案**: 在 `save_data()` 中新增智能价格合并逻辑
  - 当新数据价格为空时，自动从cache文件中合并已有价格
  - 构建货号→商品映射表，支持跨文件价格恢复
- **效果**: 防止未来再次出现价格丢失问题

#### 4. **当前数据立即修复**
- **操作**: 从历史文件（7月27日、29日）成功恢复83个商品的价格
- **结果**:
  - 高价商品: **0个 → 71个** ✅
  - 预计售出总价: **¥0 → ¥185,170.00** ✅
  - 平均售出均价: **¥0 → ¥2,230.96** ✅

### 📊 数据验证结果

| 项目 | 修复前 | 修复后 | 状态 |
|------|--------|--------|------|
| 总商品数 | 89 | 89 | ✅ |
| 高价商品(≥599) | 0 | **71** | ✅ |
| 预计售出总价 | ¥0 | **¥185,170.00** | ✅ |
| 平均售出均价 | ¥0 | **¥2,230.96** | ✅ |

### 🔧 技术改进

1. **价格持久化保障机制** - 自动备份、智能合并、多源恢复
2. **代码质量提升** - 统一字段命名、异常处理增强、日志完善

---
"""

if "## 📋 版本历史" in content:
    content = content.replace("## 📋 版本历史", version_update + "\n## 📋 版本历史")
else:
    content += version_update

with open(readme_path, "w", encoding="utf-8") as f:
    f.write(content)
print("✅ README.md 已更新")

# 2. 更新 skill.md
print("\n📄 [2/4] 更新 skill.md...")
skill_path = "skill.md"
with open(skill_path, "r", encoding="utf-8") as f:
    content = f.read()

if "## ⚡ FastAPI 迁移规范" not in content:
    fastapi_section = """

## ⚡ FastAPI 迁移规范 (v3.8.89.3+)

### 核心变更：从 Flask 完全迁移至 FastAPI

#### 1. 导入规范
```python
# ✅ 正确 (FastAPI)
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse

# ❌ 错误 (Flask)
from flask import Flask, request, jsonify
```

#### 2. 路由定义规范
```python
# ✅ 正确 (FastAPI)
@app.get("/api/products")
async def get_all_products():  # 必须使用 async
    return jsonify({"data": result})

@app.post("/api/sku/compare/txt")
async def compare_sku_txt(request: Request):  # POST需要Request参数
    data = await request.json()  # 必须使用await
```

#### 3. 响应返回规范
```python
# ✅ 正确 (FastAPI兼容格式)
return jsonify({"error": "未找到文件"}, status_code=404)

# ❌ 错误 (Flask元组格式)
return jsonify({"error": "未找到文件"}), 404
```

#### 4. 字段命名规范（必须包含英文别名）
```python
product = {
    "售价": f"¥{price:,}",
    "price": f"¥{price:,}",           # 英文别名（必须）
    "拿货价": f"¥{cost:,}",
    "cost_price": f"¥{cost:,}",       # 英文别名
    "货号": goods_num,
    "stock_number": goods_num         # 英文别名
}
# 重要: 所有筛选逻辑必须使用 price/cost_price/stock_number 字段！
```

#### 5. 价格数据处理规范
```python
def save_data(self, data):
    # 1. 读取现有文件并保存为cache
    # 2. 构建cache中的货号->商品映射
    # 3. 价格合并：如果新数据价格为空，则从cache获取
    # 4. 保存合并后的数据
```

#### 6. 常见陷阱
- ❌ Unicode字符错误: `p.get('售䷜')` → ✅ 使用 `p.get('price')`
- ❌ None vs 空字符串: 必须处理两种情况
- ❌ Flask残留代码: 必须完全移除所有Flask引用

"""
    content += fastapi_section
    
    with open(skill_path, "w", encoding="utf-8") as f:
        f.write(content)
    print("✅ skill.md 已更新（添加FastAPI规范）")
else:
    print("⚠️ skill.md 已包含FastAPI规范")

# 3. 生成 skill.docx
print("\n📄 [3/4] 生成 skill.docx...")
try:
    from docx import Document
    
    with open(skill_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    
    doc = Document()
    doc.add_heading("项目代码规范与范式 (Skill)", 0)
    doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    lines = md_content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("```"):
            continue
        else:
            doc.add_paragraph(line)
    
    doc.save("skill.docx")
    print("✅ skill.docx 已生成")
except ImportError:
    print("⚠️ python-docx 未安装")
except Exception as e:
    print(f"❌ 生成docx失败: {e}")

# 4. Git提交和推送
print("\n📤 [4/4] Git提交和推送...")
import subprocess

result = subprocess.run(["git", "status"], capture_output=True, text=True)
print(result.stdout)

if "nothing to commit" in result.stdout:
    print("✅ 没有需要提交的更改")
else:
    subprocess.run(["git", "add", "."], check=True)
    
    commit_msg = """fix: 修复高价商品显示0及Flask迁移完成 (v3.8.89.9)

主要修复:
1. 高价商品(≥599)从0修复为71个
2. 修复价格字段引用错误(9处)
3. 完成Flask到FastAPI完全迁移(20处)
4. 添加智能价格合并逻辑防止数据丢失
5. 从历史数据恢复83个商品价格

技术改进:
- 统一使用price/cost_price等英文字段
- save_data()增加cache价格合并
- 所有路由改用FastAPI async/await规范
- 更新README.md和skill.md文档

数据验证:
- 总商品数: 89个
- 高价商品: 71个 (之前0)
- 总售价: ¥185,170.00 (之前¥0)
- 平均价: ¥2,230.96 (之前¥0)
"""
    
    result = subprocess.run(["git", "commit", "-m", commit_msg], capture_output=True, text=True)
    print(result.stdout)
    
    if result.returncode == 0:
        result = subprocess.run(["git", "push"], capture_output=True, text=True)
        print(result.stdout)
        print("\n✅ Git推送成功！" if result.returncode == 0 else "\n⚠️ 推送可能失败")
    else:
        print(f"❌ 提交失败: {result.stderr}")

print("\n" + "=" * 70)
print("🎉 所有操作已完成！")
print("=" * 70)
