#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def markdown_to_docx(md_file, docx_file):
    doc = Document()

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif line.startswith('- **') and '**:' in line:
            match = re.match(r'- \*\*(.+?)\*\*:\s*(.*)', line)
            if match:
                title, text = match.groups()
                p = doc.add_paragraph(style='List Bullet')
                run_title = p.add_run(title + ': ')
                run_title.bold = True
                p.add_run(text)
            else:
                doc.add_paragraph(line, style='List Bullet')
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1

            if len(table_lines) >= 2:
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    rows.append(cells)

                if rows:
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text
                                if row_idx == 0 or (len(rows) > 1 and col_idx == 0):
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
        elif line.startswith('```'):
            lang = line[3:].strip() if len(line) > 3 else ''
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1

            if code_lines:
                p = doc.add_paragraph()
                code_text = '\n'.join(code_lines)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                run.font.size = Pt(9)
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        elif line.strip():
            if re.match(r'^\d+\.', line):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)

        i += 1

    doc.save(docx_file)
    print(f"✅ 已成功生成: {docx_file}")

if __name__ == '__main__':
    markdown_to_docx('skill.md', 'skill.docx')