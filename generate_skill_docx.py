# -*- coding: utf-8 -*-
"""
skill.md 转 skill.docx 生成工具
用于将 Markdown 格式的开发规范文档转换为 Word 格式
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path


def markdown_to_docx(md_path: str, docx_path: str):
    """
    将 Markdown 文件转换为 DOCX 格式
    
    Args:
        md_path: Markdown文件路径
        docx_path: 输出DOCX文件路径
    """
    # 读取Markdown内容
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 按行解析Markdown
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # 代码块处理
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_content = []
            else:
                # 结束代码块，写入内容
                in_code_block = False
                if code_content:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # 标题处理
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            
            # 移除Markdown格式符号
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # 粗体
            text = re.sub(r'\*([^*]+)\*', r'\1', text)      # 斜体
            
            if level == 1:
                heading = doc.add_heading(text, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                    run.font.size = Pt(24)
            elif level == 2:
                heading = doc.add_heading(text, level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                    run.font.size = Pt(18)
            elif level == 3:
                heading = doc.add_heading(text, level=2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
                    run.font.size = Pt(14)
            else:
                heading = doc.add_heading(text, level=min(level-1, 3))
                for run in heading.runs:
                    run.font.size = Pt(12)
            
            i += 1
            continue
        
        # 列表处理
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\`([^\`]+)\`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\`([^\`]+)\`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # 表格处理（简单检测）
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                if not lines[i].strip().startswith('|---'):
                    table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                # 解析表格
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]  # 移除首尾空元素
                    rows.append(cells)
                
                if len(rows) >= 1:
                    # 创建表格
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                                # 表头加粗
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
            continue
        
        # 引用块处理
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            i += 1
            continue
        
        # 分隔线
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            i += 1
            continue
        
        # 空行跳过
        if not line.strip():
            i += 1
            continue
        
        # 普通段落
        text = line.strip()
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # 粗体
        text = re.sub(r'\*([^*]+)\*', r'\1', text)      # 斜体
        text = re.sub(r'\`([^\`]+)\`', r'\1', text)      # 行内代码
        
        if text:
            p = doc.add_paragraph(text)
        
        i += 1
    
    # 保存文档
    doc.save(docx_path)
    print(f"✅ 已生成: {docx_path}")


if __name__ == '__main__':
    base_dir = Path(__file__).parent
    md_file = base_dir / 'skill.md'
    docx_file = base_dir / 'skill.docx'
    
    if md_file.exists():
        markdown_to_docx(str(md_file), str(docx_file))
        print(f"📄 文档大小: {docx_file.stat().st_size / 1024:.1f} KB")
    else:
        print(f"❌ 找不到文件: {md_file}")